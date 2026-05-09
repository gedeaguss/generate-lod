import streamlit as st
import openpyxl
import tempfile
import os
import re
import shutil
import zipfile
import io
from docx import Document

# ─────────────────────────────────────────────
#  CONFIG
# ─────────────────────────────────────────────
KOLOM_WAJIB   = ["Nomor Polis", "Nilai Klaim", "Terbilang",
                 "Akibat", "Tanggal Kejadian", "Lokasi Kejadian"]
KOLOM_NOMINAL = ["Nilai Klaim"]

st.set_page_config(
    page_title="Generate LoD Asuransi",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
#  CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700&family=DM+Mono:wght@400;500&display=swap');

html, body, [class*="css"] {
    font-family: 'Plus Jakarta Sans', sans-serif;
}

/* Warna utama */
:root {
    --primary:     #1a3a5c;
    --primary-light: #2563a8;
    --accent:      #e8a020;
    --bg:          #f4f6f9;
    --card:        #ffffff;
    --border:      #dde3ec;
    --text:        #1c2b3a;
    --muted:       #6b7c93;
    --success:     #1a7a4a;
    --warning:     #b45309;
    --error:       #b91c1c;
}

/* Header aplikasi */
.app-header {
    background: linear-gradient(135deg, #1a3a5c 0%, #2563a8 100%);
    border-radius: 16px;
    padding: 32px 40px;
    margin-bottom: 28px;
    display: flex;
    align-items: center;
    gap: 20px;
    box-shadow: 0 4px 24px rgba(26,58,92,0.18);
}
.app-header-icon { font-size: 48px; }
.app-header-title {
    color: #ffffff;
    font-size: 26px;
    font-weight: 700;
    margin: 0;
    letter-spacing: -0.3px;
}
.app-header-sub {
    color: rgba(255,255,255,0.72);
    font-size: 14px;
    margin: 4px 0 0 0;
}

/* Stepper */
.stepper {
    display: flex;
    align-items: center;
    margin-bottom: 28px;
    gap: 0;
}
.step {
    display: flex;
    align-items: center;
    gap: 10px;
    flex: 1;
}
.step-circle {
    width: 36px; height: 36px;
    border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-weight: 700; font-size: 14px;
    flex-shrink: 0;
}
.step-circle.done     { background:#1a7a4a; color:#fff; }
.step-circle.active   { background:#2563a8; color:#fff; box-shadow:0 0 0 4px rgba(37,99,168,0.18); }
.step-circle.inactive { background:#dde3ec; color:#6b7c93; }
.step-label { font-size: 13px; font-weight: 600; }
.step-label.active   { color:#2563a8; }
.step-label.done     { color:#1a7a4a; }
.step-label.inactive { color:#6b7c93; }
.step-line {
    height: 2px; flex: 1;
    margin: 0 8px;
}
.step-line.done   { background:#1a7a4a; }
.step-line.inactive { background:#dde3ec; }

/* Card */
.card {
    background: var(--card);
    border: 1px solid var(--border);
    border-radius: 12px;
    padding: 24px;
    margin-bottom: 16px;
}
.card-title {
    font-size: 15px; font-weight: 700;
    color: var(--primary);
    margin-bottom: 16px;
    display: flex; align-items: center; gap: 8px;
}

/* Badge status */
.badge {
    display: inline-flex; align-items: center; gap: 5px;
    padding: 4px 10px; border-radius: 20px;
    font-size: 12px; font-weight: 600;
}
.badge-ok      { background:#dcfce7; color:#15803d; }
.badge-warn    { background:#fef9c3; color:#a16207; }
.badge-error   { background:#fee2e2; color:#b91c1c; }
.badge-info    { background:#dbeafe; color:#1d4ed8; }

/* Tabel diagnostik */
.diag-table { width:100%; border-collapse:collapse; font-size:13px; }
.diag-table th {
    background:#f1f5f9; color:#475569;
    padding: 8px 12px; text-align:left;
    font-weight:600; border-bottom:2px solid #dde3ec;
}
.diag-table td { padding:8px 12px; border-bottom:1px solid #f1f5f9; }
.diag-table tr:last-child td { border-bottom:none; }

/* Onboarding */
.onboard-step {
    display:flex; align-items:flex-start; gap:14px;
    padding: 14px 0;
    border-bottom: 1px solid #f1f5f9;
}
.onboard-step:last-child { border-bottom:none; }
.onboard-num {
    width:32px; height:32px; border-radius:50%;
    background: linear-gradient(135deg,#1a3a5c,#2563a8);
    color:#fff; font-weight:700; font-size:14px;
    display:flex; align-items:center; justify-content:center;
    flex-shrink:0;
}
.onboard-text { font-size:14px; color:#1c2b3a; line-height:1.6; }
.onboard-text b { color:#1a3a5c; }

/* Footer */
.footer {
    text-align:center; color:#94a3b8;
    font-size:12px; margin-top:40px; padding-top:20px;
    border-top:1px solid #dde3ec;
    font-family:'DM Mono', monospace;
    letter-spacing:0.05em;
}

/* Override tombol Streamlit */
div.stButton > button {
    background: linear-gradient(135deg, #1a3a5c, #2563a8);
    color: white; border: none;
    border-radius: 8px; font-weight: 600;
    padding: 10px 24px; width: 100%;
    transition: opacity 0.2s;
}
div.stButton > button:hover { opacity: 0.88; }
div.stButton > button:disabled {
    background: #dde3ec; color: #94a3b8;
}

/* Upload area */
section[data-testid="stFileUploadDropzone"] {
    border: 2px dashed #2563a8 !important;
    border-radius: 10px !important;
    background: #f0f6ff !important;
}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  SESSION STATE
# ─────────────────────────────────────────────
def init_state():
    defaults = {
        "tahap":           1,
        "onboard_done":    False,
        "template_bytes":  None,
        "excel_bytes":     None,
        "template_name":   "",
        "excel_name":      "",
        "diag_result":     None,
        "records":         None,
        "headers":         None,
        "zip_buffer":      None,
        "generate_done":   False,
        "n_success":       0,
        "n_failed":        0,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()


# ─────────────────────────────────────────────
#  HELPER FUNCTIONS
# ─────────────────────────────────────────────
def format_rupiah(nilai):
    if nilai is None or nilai == "":
        return ""
    if isinstance(nilai, str):
        bersih = re.sub(r"[^\d,.]", "", nilai)
        bersih = bersih.replace(".", "").replace(",", ".")
        try:
            nilai = float(bersih)
        except ValueError:
            return nilai
    try:
        nilai = float(nilai)
    except (ValueError, TypeError):
        return str(nilai)
    bulat   = int(nilai)
    desimal = nilai - bulat
    ribuan  = f"{bulat:,}".replace(",", ".")
    if desimal > 0:
        sen = f"{desimal:.2f}"[1:].replace(".", ",")
        return f"Rp {ribuan}{sen}"
    return f"Rp {ribuan}"


def replace_text_in_paragraph(paragraph, replacements):
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    ada_perubahan = False
    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))
            ada_perubahan = True
    if not ada_perubahan:
        return
    paragraph.runs[0].text = full_text
    for run in paragraph.runs[1:]:
        run.text = ""


def fill_template(template_bytes, replacements):
    doc = Document(io.BytesIO(template_bytes))
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, replacements)
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_text_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            replace_text_in_paragraph(para, replacements)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def read_excel(excel_bytes):
    wb      = openpyxl.load_workbook(io.BytesIO(excel_bytes))
    ws      = wb.active
    headers = [cell.value for cell in ws[1] if cell.value]
    rows    = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(row):
            record = dict(zip(headers, row))
            for kolom in KOLOM_NOMINAL:
                if kolom in record and record[kolom] is not None:
                    record[kolom] = format_rupiah(record[kolom])
            rows.append(record)
    return headers, rows


def run_diagnostik(template_bytes, headers):
    doc        = Document(io.BytesIO(template_bytes))
    found_ok   = []
    found_split = []

    def scan(para, loc):
        full_text = "".join(r.text for r in para.runs)
        for m in re.findall(r"\{\{(.*?)\}\}", full_text):
            found_ok.append((m, loc))
        for run in para.runs:
            t = run.text
            if (("{{" in t and "}}" not in t) or
                ("}}" in t and "{{" not in t) or
                re.search(r"\{\{[^}]*$", t) or
                re.search(r"^[^{]*\}\}", t)):
                found_split.append((t.strip(), loc))

    for para in doc.paragraphs:
        scan(para, "Paragraf")
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    scan(para, f"Tabel[{ti}] Baris[{ri}] Kolom[{ci}]")
    for si, section in enumerate(doc.sections):
        for para in section.header.paragraphs:
            scan(para, f"Header[{si}]")
        for para in section.footer.paragraphs:
            scan(para, f"Footer[{si}]")

    return {
        "found_ok":    found_ok,
        "found_split": found_split,
        "headers":     headers,
    }


# ─────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    st.markdown("### 📘 Panduan Singkat")
    st.markdown("""
**Format Placeholder di Word:**
```
{{Nama Kolom}}
```
Nama di dalam kurung kurawal harus **persis sama** dengan header kolom di Excel.

---

**Kolom Wajib di Excel:**
""")
    for k in KOLOM_WAJIB:
        st.markdown(f"- `{k}`")

    st.markdown("---")
    st.markdown("""
**Tips jika placeholder tidak terbaca:**
1. Hapus placeholder di Word
2. Ketik di Notepad: `{{Nama Kolom}}`
3. Copy-paste ke Word

---

**Format output:** `.docx`
Untuk konversi ke PDF, buka file di Microsoft Word → **Save As → PDF**
""")
    st.markdown("---")
    with st.expander("❓ FAQ"):
        st.markdown("""
**Placeholder tidak terganti?**
Pastikan nama kolom Excel dan placeholder Word identik, termasuk huruf besar/kecilnya.

**Nilai klaim tidak berformat Rupiah?**
Pastikan nama kolom di Excel adalah `Nilai Klaim` (sesuai huruf besar/kecilnya).

**File gagal diproses?**
Cek kolom `Nomor Polis` — pastikan tidak ada baris yang kosong.

**Bisa tambah kolom baru?**
Bisa. Tambah kolom di Excel, lalu tambahkan `{{Nama Kolom}}` di template Word di posisi yang diinginkan.
""")


# ─────────────────────────────────────────────
#  HEADER
# ─────────────────────────────────────────────
st.markdown("""
<div class="app-header">
    <div class="app-header-icon">📄</div>
    <div>
        <p class="app-header-title">Generate LoD (Letter of Discharge) Asuransi</p>
        <p class="app-header-sub">Otomatisasi pembuatan dokumen LoD klaim asuransi secara massal</p>
    </div>
</div>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  ONBOARDING
# ─────────────────────────────────────────────
if not st.session_state.onboard_done:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">👋 Selamat Datang — Cara Menggunakan Aplikasi Ini</div>', unsafe_allow_html=True)
    st.markdown("""
<div class="onboard-step">
    <div class="onboard-num">1</div>
    <div class="onboard-text">
        <b>Siapkan template Word (.docx)</b><br>
        Buka template LoD di Microsoft Word, sesuaikan tanggal di area tanda tangan,
        lalu pastikan bagian yang ingin diisi otomatis sudah menggunakan placeholder
        dengan format <code>{{Nama Kolom}}</code>.
    </div>
</div>
<div class="onboard-step">
    <div class="onboard-num">2</div>
    <div class="onboard-text">
        <b>Siapkan file data Excel (.xlsx)</b><br>
        Pastikan baris pertama adalah nama kolom dan kolom wajib sudah tersedia:
        <b>Nomor Polis, Nilai Klaim, Terbilang, Akibat, Tanggal Kejadian, Lokasi Kejadian</b>.
        Tidak boleh ada baris kosong di antara data.
    </div>
</div>
<div class="onboard-step">
    <div class="onboard-num">3</div>
    <div class="onboard-text">
        <b>Upload kedua file & validasi</b><br>
        Aplikasi akan otomatis mengecek kelengkapan kolom dan kompatibilitas
        placeholder antara Excel dan template Word.
    </div>
</div>
<div class="onboard-step">
    <div class="onboard-num">4</div>
    <div class="onboard-text">
        <b>Generate & download</b><br>
        Klik tombol Generate — semua dokumen LoD akan dibuat sekaligus dan
        bisa didownload dalam satu file <b>.zip</b>.
    </div>
</div>
""", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    if st.button("✅ Saya Mengerti, Mulai Sekarang"):
        st.session_state.onboard_done = True
        st.rerun()
    st.stop()


# ─────────────────────────────────────────────
#  STEPPER
# ─────────────────────────────────────────────
tahap = st.session_state.tahap

def step_class(n):
    if n < tahap:  return "done"
    if n == tahap: return "active"
    return "inactive"

def line_class(n):
    return "done" if n < tahap else "inactive"

steps = ["Upload File", "Validasi & Preview", "Generate & Download"]
circles = ["✓" if i + 1 < tahap else str(i + 1) for i in range(3)]

stepper_html = '<div class="stepper">'
for i, (label, circle) in enumerate(zip(steps, circles)):
    sc = step_class(i + 1)
    stepper_html += f'''
    <div class="step">
        <div class="step-circle {sc}">{circle}</div>
        <span class="step-label {sc}">{label}</span>
    </div>'''
    if i < len(steps) - 1:
        stepper_html += f'<div class="step-line {line_class(i + 1)}"></div>'
stepper_html += "</div>"
st.markdown(stepper_html, unsafe_allow_html=True)


# ─────────────────────────────────────────────
#  TAHAP 1 — UPLOAD
# ─────────────────────────────────────────────
if tahap == 1:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">📁 Upload File</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Template Word (.docx)**")
        st.caption("Pastikan tanggal di area tanda tangan sudah disesuaikan sebelum upload.")
        f_template = st.file_uploader("Upload template", type=["docx"],
                                      label_visibility="collapsed", key="up_template")

    with col2:
        st.markdown("**Data Excel (.xlsx)**")
        st.caption("Baris pertama harus berisi nama kolom. Tidak boleh ada baris kosong di antara data.")
        f_excel = st.file_uploader("Upload Excel", type=["xlsx"],
                                   label_visibility="collapsed", key="up_excel")

    st.markdown("</div>", unsafe_allow_html=True)

    keduanya_ada = f_template is not None and f_excel is not None

    if f_template:
        st.markdown(f'<span class="badge badge-ok">✅ {f_template.name}</span>', unsafe_allow_html=True)
    if f_excel:
        st.markdown(f'<span class="badge badge-ok">✅ {f_excel.name}</span>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    col_btn, _ = st.columns([1, 2])
    with col_btn:
        if st.button("Lanjut ke Validasi →", disabled=not keduanya_ada):
            st.session_state.template_bytes = f_template.read()
            st.session_state.excel_bytes    = f_excel.read()
            st.session_state.template_name  = f_template.name
            st.session_state.excel_name     = f_excel.name

            headers, records = read_excel(st.session_state.excel_bytes)
            st.session_state.headers = headers
            st.session_state.records = records
            st.session_state.diag_result = run_diagnostik(
                st.session_state.template_bytes, headers)
            st.session_state.tahap = 2
            st.rerun()


# ─────────────────────────────────────────────
#  TAHAP 2 — VALIDASI & PREVIEW
# ─────────────────────────────────────────────
elif tahap == 2:
    headers     = st.session_state.headers
    records     = st.session_state.records
    diag        = st.session_state.diag_result
    found_ok    = diag["found_ok"]
    found_split = diag["found_split"]

    # ── Cek kolom wajib ──
    kolom_kurang = [k for k in KOLOM_WAJIB if k not in headers]
    ada_error    = len(kolom_kurang) > 0

    # ── Ringkasan file ──
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">📊 Ringkasan File</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    c1.metric("Total Data", f"{len(records)} baris")
    c2.metric("Kolom Excel", f"{len(headers)} kolom")
    c3.metric("Placeholder di Word", f"{len(found_ok)} ditemukan")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Kolom wajib ──
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">✅ Pengecekan Kolom Wajib</div>', unsafe_allow_html=True)
    if kolom_kurang:
        st.error(f"Kolom wajib berikut tidak ditemukan di Excel: **{', '.join(kolom_kurang)}**")
    else:
        st.success("Semua kolom wajib tersedia.")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Diagnostik placeholder ──
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">🔍 Diagnostik Placeholder</div>', unsafe_allow_html=True)

    if found_ok:
        ok_names = {name for name, _ in found_ok}
        rows_html = ""
        for name, loc in found_ok:
            matched = name in headers
            badge   = '<span class="badge badge-ok">✅ Matched</span>' if matched \
                      else '<span class="badge badge-error">❌ Tidak ada di Excel</span>'
            split_warn = ""
            if any(name in t for t, _ in found_split):
                split_warn = ' <span class="badge badge-warn">⚠️ Split run — diperbaiki otomatis</span>'
            rows_html += f"<tr><td><code>{{{{{name}}}}}</code></td><td>{badge}{split_warn}</td><td style='color:#94a3b8;font-size:12px'>{loc}</td></tr>"

        st.markdown(f"""
        <table class="diag-table">
            <thead><tr><th>Placeholder</th><th>Status</th><th>Lokasi</th></tr></thead>
            <tbody>{rows_html}</tbody>
        </table>""", unsafe_allow_html=True)
    else:
        st.warning("Tidak ada placeholder yang terdeteksi di template Word. Pastikan format `{{Nama Kolom}}` sudah benar.")

    st.markdown("</div>", unsafe_allow_html=True)

    # ── Preview data ──
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown('<div class="card-title">👁️ Preview Data (5 baris pertama)</div>', unsafe_allow_html=True)
    preview = records[:5]
    if preview:
        st.dataframe(preview, use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # ── Navigasi ──
    col_back, col_next, _ = st.columns([1, 1, 2])
    with col_back:
        if st.button("← Kembali"):
            st.session_state.tahap = 1
            st.rerun()
    with col_next:
        if st.button("Generate Dokumen →", disabled=ada_error):
            st.session_state.tahap = 3
            st.rerun()


# ─────────────────────────────────────────────
#  TAHAP 3 — GENERATE & DOWNLOAD
# ─────────────────────────────────────────────
elif tahap == 3:
    records       = st.session_state.records
    template_bytes = st.session_state.template_bytes
    total         = len(records)

    if not st.session_state.generate_done:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">⚙️ Generate Dokumen</div>', unsafe_allow_html=True)
        st.markdown(f"Siap membuat **{total} dokumen LoD**. Klik tombol di bawah untuk memulai.")
        st.markdown("</div>", unsafe_allow_html=True)

        col_btn, _ = st.columns([1, 2])
        with col_btn:
            if st.button("🚀 Mulai Generate"):
                progress_bar  = st.progress(0)
                status_text   = st.empty()
                zip_buf       = io.BytesIO()
                success_count = 0
                failed_list   = []

                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for idx, record in enumerate(records, start=1):
                        nomor     = idx
                        polis     = record.get("Nomor Polis", "UNKNOWN")
                        nama_file = f"LoD_{nomor:03d}_{polis}.docx"

                        status_text.markdown(
                            f'<span class="badge badge-info">⏳ [{idx}/{total}] Memproses: {nama_file}</span>',
                            unsafe_allow_html=True)
                        try:
                            docx_bytes = fill_template(template_bytes, record)
                            zf.writestr(nama_file, docx_bytes)
                            success_count += 1
                        except Exception as e:
                            failed_list.append((nama_file, str(e)))

                        progress_bar.progress(idx / total)

                status_text.empty()
                zip_buf.seek(0)

                st.session_state.zip_buffer   = zip_buf.read()
                st.session_state.generate_done = True
                st.session_state.n_success     = success_count
                st.session_state.n_failed      = len(failed_list)
                st.rerun()

    else:
        # ── Hasil ──
        n_success = st.session_state.n_success
        n_failed  = st.session_state.n_failed

        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">🎉 Generate Selesai</div>', unsafe_allow_html=True)

        c1, c2 = st.columns(2)
        c1.metric("✅ Berhasil", f"{n_success} dokumen")
        if n_failed:
            c2.metric("❌ Gagal", f"{n_failed} dokumen")

        if n_failed == 0:
            st.success(f"Semua {n_success} dokumen LoD berhasil dibuat!")
        else:
            st.warning(f"{n_success} berhasil, {n_failed} gagal. Cek kolom Nomor Polis pada baris yang bermasalah.")

        st.markdown("</div>", unsafe_allow_html=True)

        # ── Download ──
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">⬇️ Download Hasil</div>', unsafe_allow_html=True)
        st.caption("File ZIP berisi seluruh dokumen LoD dalam format .docx. "
                   "Untuk mengubah ke PDF, buka di Microsoft Word → Save As → PDF.")

        col_dl, _ = st.columns([1, 2])
        with col_dl:
            st.download_button(
                label="⬇️ Download Semua LoD (.zip)",
                data=st.session_state.zip_buffer,
                file_name="semua_LoD_klaim.zip",
                mime="application/zip",
                use_container_width=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)

        # ── Generate ulang ──
        col_back, col_reset, _ = st.columns([1, 1, 2])
        with col_back:
            if st.button("← Kembali ke Validasi"):
                st.session_state.tahap         = 2
                st.session_state.generate_done = False
                st.session_state.zip_buffer    = None
                st.rerun()
        with col_reset:
            if st.button("🔄 Generate Dokumen Baru"):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()


# ─────────────────────────────────────────────
#  FOOTER
# ─────────────────────────────────────────────
st.markdown('<div class="footer">CAPEX Management</div>', unsafe_allow_html=True)
